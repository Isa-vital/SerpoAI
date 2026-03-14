"""Convert the SRD markdown to a formatted Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import re
import os

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    hs.font.name = 'Calibri'

# --- Helper functions ---
def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h.replace('**', ''))
        run.bold = True
        run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            # Handle bold markers
            parts = re.split(r'\*\*(.+?)\*\*', cell_text)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(10)
                if j % 2 == 1:  # odd indices are bold
                    run.bold = True
    return table

def add_code_block(doc, text):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_paragraph_with_bold(doc, text, style_name=None):
    if style_name:
        p = doc.add_paragraph(style=style_name)
    else:
        p = doc.add_paragraph()
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for j, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if j % 2 == 1:
            run.bold = True
    return p

# --- Read the markdown ---
md_path = os.path.join(os.path.dirname(__file__), 'docs', 'SerpoAI_Smart_Contract_SRD.md')
with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# --- Parse and build document ---
i = 0
in_code_block = False
code_lines = []
in_table = False
table_headers = []
table_rows = []

def flush_table():
    global in_table, table_headers, table_rows
    if in_table and table_headers:
        add_table(doc, table_headers, table_rows)
        doc.add_paragraph()  # spacing
    in_table = False
    table_headers = []
    table_rows = []

while i < len(lines):
    line = lines[i].rstrip('\n')
    
    # Code block start/end
    if line.strip().startswith('```'):
        if in_code_block:
            # End code block
            flush_table()
            add_code_block(doc, '\n'.join(code_lines))
            code_lines = []
            in_code_block = False
        else:
            flush_table()
            in_code_block = True
        i += 1
        continue
    
    if in_code_block:
        code_lines.append(line)
        i += 1
        continue
    
    # Horizontal rule
    if line.strip() == '---':
        flush_table()
        i += 1
        continue
    
    # Table row
    if line.strip().startswith('|') and line.strip().endswith('|'):
        cells = [c.strip() for c in line.strip().split('|')[1:-1]]
        
        # Check if it's a separator row (e.g., |---|---|)
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            i += 1
            continue
        
        if not in_table:
            in_table = True
            table_headers = cells
        else:
            table_rows.append(cells)
        i += 1
        continue
    else:
        flush_table()
    
    # Headings
    if line.startswith('# ') and not line.startswith('## '):
        # Title - make it a title page
        title = line[2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(72)
        run = p.add_run(title)
        run.font.size = Pt(28)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        i += 1
        continue
    
    if line.startswith('## '):
        title = line[3:].strip()
        doc.add_heading(title, level=1)
        i += 1
        continue
    
    if line.startswith('### '):
        title = line[4:].strip()
        doc.add_heading(title, level=2)
        i += 1
        continue
    
    if line.startswith('#### '):
        title = line[5:].strip()
        doc.add_heading(title, level=3)
        i += 1
        continue
    
    # Empty line
    if not line.strip():
        i += 1
        continue
    
    # Blockquote
    if line.strip().startswith('> '):
        text = line.strip()[2:]
        p = add_paragraph_with_bold(doc, text)
        p.paragraph_format.left_indent = Cm(1.27)
        p.runs[0].italic = True if p.runs else None
        i += 1
        continue
    
    # Bullet point
    if line.strip().startswith('- '):
        text = line.strip()[2:]
        add_paragraph_with_bold(doc, text, 'List Bullet')
        i += 1
        continue
    
    # Numbered list
    m = re.match(r'^(\d+)\.\s+(.+)', line.strip())
    if m:
        text = m.group(2)
        add_paragraph_with_bold(doc, text, 'List Number')
        i += 1
        continue
    
    # Regular paragraph
    text = line.strip()
    if text:
        add_paragraph_with_bold(doc, text)
    
    i += 1

# Flush any remaining table
flush_table()

# --- Save ---
out_path = os.path.join(os.path.dirname(__file__), 'docs', 'SerpoAI_Smart_Contract_SRD.docx')
doc.save(out_path)
print(f"Saved: {out_path}")
